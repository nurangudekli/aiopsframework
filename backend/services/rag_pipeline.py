"""
RAG Pipeline Service.

Provides a pluggable Retrieval-Augmented Generation pipeline:
  1. Document ingestion (chunking, embedding)
  2. Vector search (via in-memory FAISS or external vector DB)
  3. Context-augmented generation
  4. File parsing (PDF, DOCX, TXT, CSV, Markdown)
  5. URL scraping for web content
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

import numpy as np

from backend.services.model_provider import call_model

logger = logging.getLogger(__name__)


# ── Data classes ────────────────────────────────────────────────
@dataclass
class Document:
    id: str
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class Chunk:
    doc_id: str
    chunk_index: int
    text: str
    embedding: Optional[np.ndarray] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class RetrievalResult:
    chunk: Chunk
    score: float


# ── In-memory vector store ──────────────────────────────────────
class InMemoryVectorStore:
    """Simple in-memory vector store backed by numpy for development / small datasets."""

    def __init__(self):
        self.chunks: List[Chunk] = []
        self._embeddings: Optional[np.ndarray] = None

    def add_chunks(self, chunks: List[Chunk]) -> None:
        self.chunks.extend(chunks)
        self._embeddings = None  # invalidate cache

    def _build_index(self):
        if self._embeddings is None and self.chunks:
            vecs = [c.embedding for c in self.chunks if c.embedding is not None]
            if vecs:
                self._embeddings = np.vstack(vecs)

    def search(self, query_embedding: np.ndarray, top_k: int = 5) -> List[RetrievalResult]:
        self._build_index()
        if self._embeddings is None or len(self._embeddings) == 0:
            return []

        # Cosine similarity
        query_norm = query_embedding / (np.linalg.norm(query_embedding) + 1e-10)
        norms = np.linalg.norm(self._embeddings, axis=1, keepdims=True) + 1e-10
        normed = self._embeddings / norms
        scores = normed @ query_norm

        top_indices = np.argsort(scores)[::-1][:top_k]
        return [
            RetrievalResult(chunk=self.chunks[i], score=float(scores[i]))
            for i in top_indices
        ]

    def clear(self) -> None:
        self.chunks = []
        self._embeddings = None


# ── Chunking ────────────────────────────────────────────────────
def chunk_text(
    text: str,
    chunk_size: int = 512,
    chunk_overlap: int = 64,
) -> List[str]:
    """Split text into overlapping chunks by character count (word-boundary aware)."""
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk_words = words[start:end]
        chunks.append(" ".join(chunk_words))
        start += chunk_size - chunk_overlap
    return chunks


# ── Embedding helper ────────────────────────────────────────────
def _get_embed_model():
    try:
        from sentence_transformers import SentenceTransformer
        return SentenceTransformer("all-MiniLM-L6-v2")
    except ImportError:
        logger.warning("sentence-transformers not installed; RAG embedding unavailable.")
        return None


def embed_texts(texts: List[str]) -> Optional[np.ndarray]:
    model = _get_embed_model()
    if model is None:
        return None
    return model.encode(texts, convert_to_numpy=True, show_progress_bar=False)


# ── RAG Pipeline ────────────────────────────────────────────────
class RAGPipeline:
    """
    End-to-end Retrieval-Augmented Generation pipeline.

    Usage:
        pipeline = RAGPipeline(vector_store=InMemoryVectorStore())
        pipeline.ingest_documents([Document(id="1", text="...")])
        answer = await pipeline.query("What is ...?", provider="azure_openai", deployment="my-deployment")
    """

    def __init__(
        self,
        vector_store: Optional[InMemoryVectorStore] = None,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        top_k: int = 5,
    ):
        self.vector_store = vector_store or InMemoryVectorStore()
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.top_k = top_k

    def ingest_documents(self, documents: List[Document]) -> int:
        """Chunk, embed, and index documents. Returns number of chunks created."""
        all_chunks: List[Chunk] = []

        for doc in documents:
            texts = chunk_text(doc.text, self.chunk_size, self.chunk_overlap)
            for idx, t in enumerate(texts):
                all_chunks.append(Chunk(doc_id=doc.id, chunk_index=idx, text=t, metadata=doc.metadata))

        # Embed all chunks
        chunk_texts = [c.text for c in all_chunks]
        embeddings = embed_texts(chunk_texts)
        if embeddings is not None:
            for i, chunk in enumerate(all_chunks):
                chunk.embedding = embeddings[i]

        self.vector_store.add_chunks(all_chunks)
        logger.info("Ingested %d documents into %d chunks", len(documents), len(all_chunks))
        return len(all_chunks)

    def retrieve(self, query: str) -> List[RetrievalResult]:
        """Retrieve top-k relevant chunks for a query."""
        query_emb = embed_texts([query])
        if query_emb is None:
            return []
        return self.vector_store.search(query_emb[0], top_k=self.top_k)

    async def query(
        self,
        question: str,
        provider: str,
        deployment: str,
        system_message: Optional[str] = None,
        **model_params: Any,
    ) -> Dict[str, Any]:
        """Full RAG: retrieve context, augment prompt, call model."""
        retrieved = self.retrieve(question)
        context_parts = [r.chunk.text for r in retrieved]
        context = "\n\n---\n\n".join(context_parts)

        augmented_prompt = (
            f"Use the following context to answer the question. "
            f"If the context does not contain enough information, say so.\n\n"
            f"Context:\n{context}\n\n"
            f"Question: {question}"
        )

        messages = []
        if system_message:
            messages.append({"role": "system", "content": system_message})
        messages.append({"role": "user", "content": augmented_prompt})

        resp = await call_model(provider, deployment, messages, **model_params)
        return {
            "answer": resp.text,
            "latency_ms": resp.latency_ms,
            "tokens_prompt": resp.tokens_prompt,
            "tokens_completion": resp.tokens_completion,
            "retrieved_chunks": [
                {"text": r.chunk.text, "score": round(r.score, 4), "doc_id": r.chunk.doc_id}
                for r in retrieved
            ],
            "error": resp.error,
        }

    def get_stats(self) -> Dict[str, Any]:
        """Return stats about the current vector store."""
        chunks = self.vector_store.chunks
        doc_ids = set(c.doc_id for c in chunks)
        total_chars = sum(len(c.text) for c in chunks)
        return {
            "total_chunks": len(chunks),
            "total_documents": len(doc_ids),
            "document_ids": sorted(doc_ids),
            "total_characters": total_chars,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "top_k": self.top_k,
        }


# ── File parsing helpers ───────────────────────────────────────
def parse_file_to_text(filename: str, content: bytes) -> str:
    """Extract plain text from various file formats."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "txt" or ext == "md":
        return content.decode("utf-8", errors="replace")

    if ext == "csv":
        import csv
        reader = csv.reader(io.StringIO(content.decode("utf-8", errors="replace")))
        rows = list(reader)
        return "\n".join(", ".join(row) for row in rows)

    if ext == "json" or ext == "jsonl":
        import json
        text_parts: List[str] = []
        if ext == "jsonl":
            for line in content.decode("utf-8", errors="replace").strip().split("\n"):
                if line.strip():
                    obj = json.loads(line)
                    text_parts.append(json.dumps(obj, indent=2))
        else:
            data = json.loads(content.decode("utf-8", errors="replace"))
            if isinstance(data, list):
                for item in data:
                    text_parts.append(json.dumps(item, indent=2))
            else:
                text_parts.append(json.dumps(data, indent=2))
        return "\n\n".join(text_parts)

    if ext == "pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=content, filetype="pdf")
            pages = [page.get_text() for page in doc]
            doc.close()
            return "\n\n".join(pages)
        except ImportError:
            logger.warning("PyMuPDF (fitz) not installed; PDF parsing unavailable.")
            return content.decode("utf-8", errors="replace")

    if ext in ("docx",):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            logger.warning("python-docx not installed; DOCX parsing unavailable.")
            return content.decode("utf-8", errors="replace")

    # Fallback: treat as plain text
    return content.decode("utf-8", errors="replace")


# ── URL scraping helper ────────────────────────────────────────
async def scrape_url(url: str) -> str:
    """Fetch a URL and extract readable text content."""
    import httpx

    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.get(url, follow_redirects=True)
        resp.raise_for_status()
        content_type = resp.headers.get("content-type", "")

        if "application/json" in content_type:
            import json
            return json.dumps(resp.json(), indent=2)

        html = resp.text

    # Simple HTML → text extraction (strip tags)
    try:
        from html.parser import HTMLParser

        class _TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self._parts: List[str] = []
                self._skip = False

            def handle_starttag(self, tag, attrs):
                if tag in ("script", "style", "nav", "footer", "header"):
                    self._skip = True

            def handle_endtag(self, tag):
                if tag in ("script", "style", "nav", "footer", "header"):
                    self._skip = False
                if tag in ("p", "div", "br", "h1", "h2", "h3", "h4", "li", "tr"):
                    self._parts.append("\n")

            def handle_data(self, data):
                if not self._skip:
                    cleaned = data.strip()
                    if cleaned:
                        self._parts.append(cleaned)

        extractor = _TextExtractor()
        extractor.feed(html)
        text = " ".join(extractor._parts)
        # Collapse whitespace
        text = re.sub(r"\s+", " ", text).strip()
        return text
    except Exception:
        # Fallback: regex strip tags
        return re.sub(r"<[^>]+>", " ", html)
